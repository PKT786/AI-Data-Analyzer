"""
=========================================================
AI Data Analyzer Pro
Enterprise Report Builder

Author : Punit Tech Hub
Version : 3.0
=========================================================
"""

from __future__ import annotations

from datetime import datetime
from typing import Any, Dict, List

import pandas as pd
import numpy as np

from utils.logger import get_logger
from utils.validators import validate_dataframe

logger = get_logger(__name__)


# =========================================================
# REPORT BUILDER
# =========================================================

class ReportBuilder:
    """
    Enterprise Report Builder

    Generates

    ✓ Executive Summary

    ✓ Dataset Overview

    ✓ Data Quality Report

    ✓ Business Insights

    ✓ KPI Summary

    ✓ Recommendations

    ✓ Exportable Report

    Every report is returned as a dictionary.
    """

    # =====================================================
    # INITIALIZATION
    # =====================================================

    def __init__(self):

        logger.info(
            "ReportBuilder initialized."
        )

    # =====================================================
    # INTERNAL COPY
    # =====================================================

    def _copy_df(
        self,
        df: pd.DataFrame,
    ) -> pd.DataFrame:

        validate_dataframe(df)

        return df.copy(deep=True)

    # =====================================================
    # REPORT METADATA
    # =====================================================

    def report_metadata(
        self,
    ) -> Dict[str, Any]:

        return {

            "title":
                "AI Data Analyzer Report",

            "version":
                "3.0 Enterprise",

            "generated_on":
                datetime.now().strftime(
                    "%d-%m-%Y %H:%M:%S"
                ),

            "generated_by":
                "AI Data Analyzer Pro",

        }

    # =====================================================
    # DATASET OVERVIEW
    # =====================================================

    def dataset_overview(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        numeric_columns = working_df.select_dtypes(

            include=np.number

        ).columns.tolist()

        categorical_columns = working_df.select_dtypes(

            include=[
                "object",
                "category",
                "string",
            ]

        ).columns.tolist()

        datetime_columns = working_df.select_dtypes(

            include=[
                "datetime",
                "datetimetz",
            ]

        ).columns.tolist()

        return {

            "rows":
                len(working_df),

            "columns":
                len(working_df.columns),

            "numeric_columns":
                len(numeric_columns),

            "categorical_columns":
                len(categorical_columns),

            "datetime_columns":
                len(datetime_columns),

            "memory_mb":
                round(

                    working_df.memory_usage(
                        deep=True
                    ).sum()

                    / 1024
                    / 1024,

                    2,

                ),

            "duplicates":
                int(
                    working_df.duplicated().sum()
                ),

            "missing_cells":
                int(
                    working_df.isna().sum().sum()
                ),

        }

    # =====================================================
    # EXECUTIVE SUMMARY
    # =====================================================

    def executive_summary(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        overview = self.dataset_overview(df)

        summary_text = (

            f"The uploaded dataset contains "

            f"{overview['rows']} rows and "

            f"{overview['columns']} columns. "

            f"It includes "

            f"{overview['numeric_columns']} numeric columns, "

            f"{overview['categorical_columns']} categorical columns "

            f"and "

            f"{overview['datetime_columns']} datetime columns. "

            f"There are "

            f"{overview['duplicates']} duplicate rows "

            f"and "

            f"{overview['missing_cells']} missing values."

        )

        return {

            "summary_text":
                summary_text,

            "dataset_overview":
                overview,

        }

    # =====================================================
    # BASIC REPORT
    # =====================================================

    def basic_report(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        report = {

            "metadata":
                self.report_metadata(),

            "executive_summary":
                self.executive_summary(df),

        }

        logger.info(
            "Basic report generated."
        )

        return report
        # =====================================================
    # DATA QUALITY REPORT
    # =====================================================

    def data_quality_report(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        total_cells = (
            working_df.shape[0]
            * working_df.shape[1]
        )

        missing_cells = int(
            working_df.isna().sum().sum()
        )

        duplicate_rows = int(
            working_df.duplicated().sum()
        )

        return {

            "total_rows":
                len(working_df),

            "total_columns":
                len(working_df.columns),

            "missing_cells":
                missing_cells,

            "missing_percentage":
                round(

                    (missing_cells / total_cells) * 100,

                    2,

                ) if total_cells else 0,

            "duplicate_rows":
                duplicate_rows,

            "duplicate_percentage":
                round(

                    duplicate_rows

                    / len(working_df)

                    * 100,

                    2,

                ) if len(working_df) else 0,

        }

    # =====================================================
    # MISSING VALUE REPORT
    # =====================================================

    def missing_value_report(
        self,
        df: pd.DataFrame,
    ) -> pd.DataFrame:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        report = pd.DataFrame({

            "Column":
                working_df.columns,

            "Missing Count":
                working_df.isna().sum().values,

            "Missing Percentage":
                (

                    working_df.isna().mean()

                    * 100

                ).round(2).values,

        })

        report = report.sort_values(

            "Missing Count",

            ascending=False,

        )

        return report

    # =====================================================
    # DUPLICATE REPORT
    # =====================================================

    def duplicate_report(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        duplicate_rows = working_df[
            working_df.duplicated()
        ]

        return {

            "duplicate_count":
                len(duplicate_rows),

            "duplicate_percentage":
                round(

                    len(duplicate_rows)

                    / len(working_df)

                    * 100,

                    2,

                ) if len(working_df) else 0,

            "duplicates":
                duplicate_rows,

        }

    # =====================================================
    # DATA TYPE SUMMARY
    # =====================================================

    def datatype_summary(
        self,
        df: pd.DataFrame,
    ) -> pd.DataFrame:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        summary = pd.DataFrame({

            "Column":
                working_df.columns,

            "Data Type":
                working_df.dtypes.astype(str).values,

            "Missing":
                working_df.isna().sum().values,

            "Unique":
                working_df.nunique().values,

        })

        return summary

    # =====================================================
    # NUMERIC SUMMARY
    # =====================================================

    def numeric_summary(
        self,
        df: pd.DataFrame,
    ) -> pd.DataFrame:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        numeric_df = working_df.select_dtypes(

            include=np.number

        )

        if numeric_df.empty:

            return pd.DataFrame()

        summary = numeric_df.describe().T

        summary["missing"] = (

            numeric_df.isna().sum()

        )

        summary["variance"] = (

            numeric_df.var()

        )

        summary["skewness"] = (

            numeric_df.skew()

        )

        summary["kurtosis"] = (

            numeric_df.kurt()

        )

        return summary.round(2)

    # =====================================================
    # CATEGORICAL SUMMARY
    # =====================================================

    def categorical_summary(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Dict[str, Any]]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        categorical = working_df.select_dtypes(

            include=[

                "object",

                "category",

                "string",

            ]

        )

        results = {}

        for column in categorical.columns:

            value_counts = categorical[column].value_counts(

                dropna=False

            )

            results[column] = {

                "unique_values":

                    int(

                        categorical[column].nunique()

                    ),

                "top_value":

                    value_counts.index[0]

                    if len(value_counts)

                    else None,

                "top_count":

                    int(value_counts.iloc[0])

                    if len(value_counts)

                    else 0,

            }

        return results
        # =====================================================
    # CORRELATION REPORT
    # =====================================================

    def correlation_report(
        self,
        df: pd.DataFrame,
    ) -> pd.DataFrame:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        numeric_df = working_df.select_dtypes(
            include=np.number
        )

        if numeric_df.shape[1] < 2:
            return pd.DataFrame()

        corr = numeric_df.corr()

        return corr.round(3)

    # =====================================================
    # DISTRIBUTION SUMMARY
    # =====================================================

    def distribution_summary(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Dict[str, Any]]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        numeric_df = working_df.select_dtypes(
            include=np.number
        )

        results = {}

        for column in numeric_df.columns:

            results[column] = {

                "minimum":
                    float(numeric_df[column].min()),

                "maximum":
                    float(numeric_df[column].max()),

                "mean":
                    float(numeric_df[column].mean()),

                "median":
                    float(numeric_df[column].median()),

                "std":
                    float(numeric_df[column].std()),

                "variance":
                    float(numeric_df[column].var()),

                "skewness":
                    float(numeric_df[column].skew()),

                "kurtosis":
                    float(numeric_df[column].kurt()),

            }

        return results

    # =====================================================
    # OUTLIER SUMMARY
    # =====================================================

    def outlier_summary(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, int]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        numeric_df = working_df.select_dtypes(
            include=np.number
        )

        results = {}

        for column in numeric_df.columns:

            q1 = numeric_df[column].quantile(0.25)

            q3 = numeric_df[column].quantile(0.75)

            iqr = q3 - q1

            lower = q1 - (1.5 * iqr)

            upper = q3 + (1.5 * iqr)

            outliers = numeric_df[
                (numeric_df[column] < lower)
                |
                (numeric_df[column] > upper)
            ]

            results[column] = len(outliers)

        return results

    # =====================================================
    # COLUMN STATISTICS
    # =====================================================

    def column_statistics(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        results = {}

        for column in working_df.columns:

            results[column] = {

                "dtype":
                    str(working_df[column].dtype),

                "unique":
                    int(
                        working_df[column].nunique()
                    ),

                "missing":
                    int(
                        working_df[column].isna().sum()
                    ),

                "memory_bytes":
                    int(
                        working_df[column]
                        .memory_usage(
                            deep=True
                        )
                    ),

            }

        return results

    # =====================================================
    # DATASET HEALTH SCORE
    # =====================================================

    def dataset_health_score(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        total_cells = (
            working_df.shape[0]
            * working_df.shape[1]
        )

        missing = int(
            working_df.isna().sum().sum()
        )

        duplicates = int(
            working_df.duplicated().sum()
        )

        score = 100

        if total_cells > 0:

            score -= (
                missing
                / total_cells
            ) * 40

        if len(working_df) > 0:

            score -= (
                duplicates
                / len(working_df)
            ) * 30

        score = max(0, min(100, round(score, 2)))

        if score >= 90:
            status = "Excellent"

        elif score >= 75:
            status = "Good"

        elif score >= 60:
            status = "Average"

        else:
            status = "Poor"

        return {

            "score": score,

            "status": status,

        }

    # =====================================================
    # STATISTICAL OVERVIEW
    # =====================================================

    def statistical_overview(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        return {

            "numeric_summary":
                self.numeric_summary(df),

            "categorical_summary":
                self.categorical_summary(df),

            "distribution":
                self.distribution_summary(df),

            "outliers":
                self.outlier_summary(df),

            "health":
                self.dataset_health_score(df),

        }
        # =====================================================
    # KPI SUMMARY
    # =====================================================

    def kpi_summary(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        numeric = working_df.select_dtypes(
            include=np.number
        )

        return {

            "total_records":
                len(working_df),

            "total_columns":
                len(working_df.columns),

            "numeric_columns":
                len(numeric.columns),

            "missing_values":
                int(
                    working_df.isna().sum().sum()
                ),

            "duplicate_rows":
                int(
                    working_df.duplicated().sum()
                ),

            "memory_mb":
                round(

                    working_df.memory_usage(
                        deep=True
                    ).sum()

                    / 1024
                    / 1024,

                    2,

                ),

        }

    # =====================================================
    # BUSINESS INSIGHTS
    # =====================================================

    def business_insights(
        self,
        df: pd.DataFrame,
    ) -> List[str]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        insights = []

        if working_df.isna().sum().sum() > 0:

            insights.append(
                "Dataset contains missing values that should be cleaned before modelling."
            )

        if working_df.duplicated().sum() > 0:

            insights.append(
                "Duplicate records detected. Removing duplicates may improve data quality."
            )

        numeric = working_df.select_dtypes(
            include=np.number
        )

        if len(numeric.columns) > 0:

            insights.append(
                f"Dataset contains {len(numeric.columns)} numeric features suitable for machine learning."
            )

        categorical = working_df.select_dtypes(
            include=["object", "category"]
        )

        if len(categorical.columns) > 0:

            insights.append(
                f"{len(categorical.columns)} categorical columns may require encoding."
            )

        if not insights:

            insights.append(
                "Dataset appears clean and ready for analysis."
            )

        return insights

    # =====================================================
    # RISK ANALYSIS
    # =====================================================

    def risk_analysis(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, str]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        risks = {}

        missing = working_df.isna().sum().sum()

        duplicates = working_df.duplicated().sum()

        if missing > 0:

            risks["Missing Values"] = "Medium"

        else:

            risks["Missing Values"] = "Low"

        if duplicates > 0:

            risks["Duplicate Records"] = "Medium"

        else:

            risks["Duplicate Records"] = "Low"

        numeric = working_df.select_dtypes(
            include=np.number
        )

        if len(numeric.columns) == 0:

            risks["Machine Learning Readiness"] = "High"

        else:

            risks["Machine Learning Readiness"] = "Low"

        return risks

    # =====================================================
    # DATA QUALITY RECOMMENDATIONS
    # =====================================================

    def recommendations(
        self,
        df: pd.DataFrame,
    ) -> List[str]:

        validate_dataframe(df)

        working_df = self._copy_df(df)

        rec = []

        if working_df.isna().sum().sum():

            rec.append(
                "Handle missing values using Mean, Median or Mode."
            )

        if working_df.duplicated().sum():

            rec.append(
                "Remove duplicate rows."
            )

        numeric = working_df.select_dtypes(
            include=np.number
        )

        if len(numeric.columns):

            rec.append(
                "Check numeric columns for outliers."
            )

        categorical = working_df.select_dtypes(
            include=["object", "category"]
        )

        if len(categorical.columns):

            rec.append(
                "Encode categorical variables before machine learning."
            )

        rec.append(
            "Validate data before building dashboards."
        )

        return rec

    # =====================================================
    # EXECUTIVE RECOMMENDATION
    # =====================================================

    def executive_recommendation(
        self,
        df: pd.DataFrame,
    ) -> str:

        score = self.dataset_health_score(df)

        if score["score"] >= 90:

            return (
                "Dataset is production ready."
            )

        if score["score"] >= 75:

            return (
                "Minor cleaning recommended before analytics."
            )

        if score["score"] >= 60:

            return (
                "Moderate data cleaning required before reporting."
            )

        return (
            "Extensive data cleaning is recommended before analysis."
        )

    # =====================================================
    # ACTION ITEMS
    # =====================================================

    def action_items(
        self,
        df: pd.DataFrame,
    ) -> List[str]:

        return [

            "Review missing values",

            "Remove duplicates",

            "Detect outliers",

            "Validate data types",

            "Generate dashboard",

            "Generate AI insights",

            "Export cleaned dataset",

        ]

    # =====================================================
    # FINAL CONCLUSION
    # =====================================================

    def final_conclusion(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:

        return {

            "health":
                self.dataset_health_score(df),

            "recommendation":
                self.executive_recommendation(df),

            "business_insights":
                self.business_insights(df),

            "action_items":
                self.action_items(df),

        }
        # =====================================================
    # BUILD COMPLETE REPORT
    # =====================================================

    def build_complete_report(
        self,
        df: pd.DataFrame,
    ) -> Dict[str, Any]:
        """
        Generate complete enterprise report.
        """

        validate_dataframe(df)

        working_df = self._copy_df(df)

        report = {

            "metadata":
                self.report_metadata(),

            "executive_summary":
                self.executive_summary(
                    working_df
                ),

            "dataset_overview":
                self.dataset_overview(
                    working_df
                ),

            "data_quality":
                self.data_quality_report(
                    working_df
                ),

            "datatype_summary":
                self.datatype_summary(
                    working_df
                ),

            "numeric_summary":
                self.numeric_summary(
                    working_df
                ),

            "categorical_summary":
                self.categorical_summary(
                    working_df
                ),

            "correlation":
                self.correlation_report(
                    working_df
                ),

            "statistics":
                self.statistical_overview(
                    working_df
                ),

            "kpi_summary":
                self.kpi_summary(
                    working_df
                ),

            "risk_analysis":
                self.risk_analysis(
                    working_df
                ),

            "recommendations":
                self.recommendations(
                    working_df
                ),

            "business_insights":
                self.business_insights(
                    working_df
                ),

            "conclusion":
                self.final_conclusion(
                    working_df
                ),

        }

        logger.info(
            "Complete report generated."
        )

        return report

    # =====================================================
    # EXPORT JSON
    # =====================================================

    def export_json(
        self,
        report: Dict[str, Any],
    ) -> Dict[str, Any]:
        """
        JSON ready report.
        """

        return report

    # =====================================================
    # EXPORT MARKDOWN
    # =====================================================

    def export_markdown(
        self,
        report: Dict[str, Any],
    ) -> str:

        md = "# AI Data Analyzer Report\n\n"

        md += "## Executive Summary\n\n"

        md += report["executive_summary"][
            "summary_text"
        ]

        md += "\n\n"

        md += "## Dataset Overview\n\n"

        overview = report["dataset_overview"]

        for key, value in overview.items():

            md += f"- **{key}** : {value}\n"

        md += "\n"

        md += "## Recommendations\n\n"

        for item in report["recommendations"]:

            md += f"- {item}\n"

        return md

    # =====================================================
    # EXPORT HTML
    # =====================================================

    def export_html(
        self,
        report: Dict[str, Any],
    ) -> str:
        """Premium, comprehensive HTML report covering every section of
        the report dict (not just the executive summary)."""

        meta = report.get("metadata", {})
        summary = report.get("executive_summary", {}).get("summary_text", "")
        overview = report.get("dataset_overview", {})
        quality = report.get("data_quality", {})
        kpis = report.get("kpi_summary", {})
        insights = report.get("business_insights", [])
        recommendations = report.get("recommendations", [])
        conclusion = report.get("conclusion", {})
        health = conclusion.get("health", {})

        def kpi_card(label, value):
            return (
                f'<div class="kpi-card"><div class="kpi-label">{label}</div>'
                f'<div class="kpi-value">{value}</div></div>'
            )

        kpi_html = "".join([
            kpi_card("Total Rows", overview.get("rows", "N/A")),
            kpi_card("Total Columns", overview.get("columns", "N/A")),
            kpi_card("Missing Cells", quality.get("missing_cells", "N/A")),
            kpi_card("Duplicate Rows", quality.get("duplicate_rows", "N/A")),
            kpi_card("Data Health", f"{health.get('score', 'N/A')} ({health.get('status', '')})"),
        ])

        overview_rows = "".join(
            f"<tr><td>{k.replace('_', ' ').title()}</td><td>{v}</td></tr>"
            for k, v in overview.items()
        )

        insights_html = "".join(f"<li>{i}</li>" for i in insights)
        recs_html = "".join(f"<li>{r}</li>" for r in recommendations)

        return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>AI Data Analyzer Report</title>
<style>
    body {{
        margin: 0; font-family: 'Segoe UI', Arial, sans-serif;
        background: #F5F7FB; color: #0F172A;
    }}
    .header {{
        background: linear-gradient(90deg, #1C2D60, #483A96);
        color: #F5F0E6; padding: 34px 40px;
    }}
    .header h1 {{ margin: 0; font-size: 28px; }}
    .header p {{ margin: 6px 0 0; opacity: 0.85; font-size: 13px; }}
    .kpi-row {{
        display: flex; flex-wrap: wrap; gap: 14px; padding: 24px 40px 0;
    }}
    .kpi-card {{
        background: #FFFFFF; border-radius: 12px; padding: 16px 18px;
        min-width: 150px; flex: 1; box-shadow: 0 4px 14px rgba(13,27,42,0.08);
    }}
    .kpi-label {{ color: #64748B; font-size: 12px; text-transform: uppercase; }}
    .kpi-value {{ font-size: 22px; font-weight: 700; color: #1C2D60; }}
    .section {{
        background: #FFFFFF; border-radius: 14px; padding: 22px 26px;
        margin: 18px 40px; box-shadow: 0 4px 14px rgba(13,27,42,0.06);
    }}
    .section h2 {{ margin-top: 0; font-size: 18px; color: #1C2D60; }}
    table {{ width: 100%; border-collapse: collapse; }}
    td {{ padding: 6px 4px; border-bottom: 1px solid #E7EAF0; font-size: 14px; }}
    ul {{ margin: 0; padding-left: 20px; }}
    li {{ margin-bottom: 6px; font-size: 14px; }}
    .footer {{ text-align: center; color: #94A3B8; font-size: 12px; padding: 24px; }}
</style>
</head>
<body>
    <div class="header">
        <h1>📊 AI Data Analyzer Report</h1>
        <p>{meta.get('generated_by', 'AI Data Analyzer Pro')} &bull; {meta.get('generated_on', '')} &bull; v{meta.get('version', '')}</p>
    </div>
    <div class="kpi-row">{kpi_html}</div>
    <div class="section">
        <h2>Executive Summary</h2>
        <p>{summary}</p>
    </div>
    <div class="section">
        <h2>Dataset Overview</h2>
        <table>{overview_rows}</table>
    </div>
    <div class="section">
        <h2>Business Insights</h2>
        <ul>{insights_html}</ul>
    </div>
    <div class="section">
        <h2>Recommendations</h2>
        <ul>{recs_html}</ul>
    </div>
    <div class="footer">© 2026 Punit Tech Hub &bull; AI Data Analyzer Pro</div>
</body>
</html>"""

    # =====================================================
    # EXPORT PDF
    # =====================================================

    def export_pdf(
        self,
        report: Dict[str, Any],
    ) -> bytes:
        """Render the report to a clean, multi-section PDF."""

        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.pdfgen import canvas as pdf_canvas

        meta = report.get("metadata", {})
        summary = report.get("executive_summary", {}).get("summary_text", "")
        overview = report.get("dataset_overview", {})
        insights = report.get("business_insights", [])
        recommendations = report.get("recommendations", [])
        conclusion = report.get("conclusion", {})
        health = conclusion.get("health", {})

        buffer = io.BytesIO()
        c = pdf_canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        primary = HexColor("#1C2D60")
        text_color = HexColor("#0F172A")
        muted = HexColor("#64748B")

        def header(title):
            c.setFillColor(primary)
            c.rect(0, height - 2.4 * cm, width, 2.4 * cm, fill=True, stroke=False)
            c.setFillColor(HexColor("#FFFFFF"))
            c.setFont("Helvetica-Bold", 16)
            c.drawString(1.5 * cm, height - 1.6 * cm, title)

        def wrapped_text(text, x, y, max_width_chars=95, font="Helvetica", size=10, leading=14, color=None):
            c.setFont(font, size)
            c.setFillColor(color or text_color)
            import textwrap
            for line in textwrap.wrap(text, max_width_chars):
                c.drawString(x, y, line)
                y -= leading
            return y

        header("AI Data Analyzer Report")
        y = height - 3.2 * cm
        c.setFont("Helvetica", 9)
        c.setFillColor(muted)
        c.drawString(1.5 * cm, y, f"Generated {meta.get('generated_on', '')} by {meta.get('generated_by', '')}")
        y -= 0.9 * cm

        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(primary)
        c.drawString(1.5 * cm, y, "Executive Summary")
        y -= 0.6 * cm
        y = wrapped_text(summary, 1.5 * cm, y)
        y -= 0.5 * cm

        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(primary)
        c.drawString(1.5 * cm, y, f"Data Health Score: {health.get('score', 'N/A')} ({health.get('status', '')})")
        y -= 0.8 * cm

        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(primary)
        c.drawString(1.5 * cm, y, "Dataset Overview")
        y -= 0.6 * cm
        for k, v in overview.items():
            if y < 3 * cm:
                c.showPage()
                header("AI Data Analyzer Report")
                y = height - 3.2 * cm
            c.setFont("Helvetica", 10)
            c.setFillColor(text_color)
            c.drawString(1.7 * cm, y, f"{k.replace('_', ' ').title()}: {v}")
            y -= 0.5 * cm

        y -= 0.4 * cm
        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(primary)
        c.drawString(1.5 * cm, y, "Business Insights")
        y -= 0.6 * cm
        for item in insights:
            if y < 3 * cm:
                c.showPage()
                header("AI Data Analyzer Report")
                y = height - 3.2 * cm
            y = wrapped_text(f"\u2022 {item}", 1.7 * cm, y)

        y -= 0.4 * cm
        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(primary)
        if y < 3 * cm:
            c.showPage()
            header("AI Data Analyzer Report")
            y = height - 3.2 * cm
        c.drawString(1.5 * cm, y, "Recommendations")
        y -= 0.6 * cm
        for item in recommendations:
            if y < 3 * cm:
                c.showPage()
                header("AI Data Analyzer Report")
                y = height - 3.2 * cm
            y = wrapped_text(f"\u2022 {item}", 1.7 * cm, y)

        c.save()
        return buffer.getvalue()

    # =====================================================
    # EXPORT WORD
    # =====================================================

    def export_word(
        self,
        report: Dict[str, Any],
    ) -> bytes:
        """Render the report to a structured Word (.docx) document."""

        import io
        from docx import Document
        from docx.shared import Pt, RGBColor

        meta = report.get("metadata", {})
        summary = report.get("executive_summary", {}).get("summary_text", "")
        overview = report.get("dataset_overview", {})
        quality = report.get("data_quality", {})
        insights = report.get("business_insights", [])
        recommendations = report.get("recommendations", [])
        conclusion = report.get("conclusion", {})
        health = conclusion.get("health", {})

        doc = Document()

        title = doc.add_heading("AI Data Analyzer Report", level=0)
        title.runs[0].font.color.rgb = RGBColor(0x1C, 0x2D, 0x60)

        meta_p = doc.add_paragraph(
            f"Generated {meta.get('generated_on', '')} by {meta.get('generated_by', '')}"
        )
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(summary)

        doc.add_heading(f"Data Health Score: {health.get('score', 'N/A')} ({health.get('status', '')})", level=2)

        doc.add_heading("Dataset Overview", level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for k, v in overview.items():
            row = table.add_row().cells
            row[0].text = k.replace("_", " ").title()
            row[1].text = str(v)

        doc.add_heading("Data Quality", level=1)
        table2 = doc.add_table(rows=0, cols=2)
        table2.style = "Light Grid Accent 1"
        for k, v in quality.items():
            row = table2.add_row().cells
            row[0].text = k.replace("_", " ").title()
            row[1].text = str(v)

        doc.add_heading("Business Insights", level=1)
        for item in insights:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Recommendations", level=1)
        for item in recommendations:
            doc.add_paragraph(item, style="List Bullet")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # =====================================================
    # EXPORT EXCEL
    # =====================================================

    def export_excel(
        self,
        report: Dict[str, Any],
    ) -> bytes:
        """Render the report to a multi-sheet Excel workbook."""

        import io
        import pandas as _pd

        overview = report.get("dataset_overview", {})
        quality = report.get("data_quality", {})
        kpis = report.get("kpi_summary", {})
        insights = report.get("business_insights", [])
        recommendations = report.get("recommendations", [])
        conclusion = report.get("conclusion", {})
        health = conclusion.get("health", {})

        buffer = io.BytesIO()
        with _pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:

            _pd.DataFrame(
                [{"Metric": "Data Health Score", "Value": health.get("score")},
                 {"Metric": "Health Status", "Value": health.get("status")}]
                + [{"Metric": k.replace("_", " ").title(), "Value": v} for k, v in overview.items()]
            ).to_excel(writer, sheet_name="Overview", index=False)

            _pd.DataFrame(
                [{"Metric": k.replace("_", " ").title(), "Value": v} for k, v in quality.items()]
            ).to_excel(writer, sheet_name="Data Quality", index=False)

            _pd.DataFrame(
                [{"KPI": k.replace("_", " ").title(), "Value": v} for k, v in kpis.items()]
            ).to_excel(writer, sheet_name="KPIs", index=False)

            _pd.DataFrame({"Business Insights": insights}).to_excel(
                writer, sheet_name="Insights", index=False
            )
            _pd.DataFrame({"Recommendations": recommendations}).to_excel(
                writer, sheet_name="Recommendations", index=False
            )

        return buffer.getvalue()

    # =====================================================
    # SIMPLE TEXT REPORT
    # =====================================================

    def export_text(
        self,
        report: Dict[str, Any],
    ) -> str:

        text = ""

        text += "AI DATA ANALYZER REPORT\n"

        text += "=" * 50 + "\n\n"

        text += report["executive_summary"][
            "summary_text"
        ]

        text += "\n\n"

        text += "Recommendations\n"

        text += "-" * 30 + "\n"

        for item in report["recommendations"]:

            text += f"• {item}\n"

        return text

    # =====================================================
    # REPORT INFO
    # =====================================================

    @property
    def version(self):

        return "3.0 Enterprise"

    def __repr__(self):

        return (

            f"ReportBuilder("

            f"version='{self.version}')"

        )


# =====================================================
# GLOBAL INSTANCE
# =====================================================

report_builder = ReportBuilder()
